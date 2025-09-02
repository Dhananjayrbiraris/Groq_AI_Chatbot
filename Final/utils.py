import os
from groq import Groq
from dotenv import load_dotenv
import requests
from urllib.parse import urlparse
from bs4 import BeautifulSoup
import re
import json
from typing import List, Dict, Any, Optional
import numpy as np
import pandas as pd
import csv
import docx
from PyPDF2 import PdfReader
import logging

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# LangChain imports
from langchain.text_splitter import RecursiveCharacterTextSplitter
#from langchain.embeddings import HuggingFaceEmbeddings
from langchain_huggingface import HuggingFaceEmbeddings  # ✅

from langchain_chroma import Chroma
from langchain.schema import Document
from langchain.chains import RetrievalQA
from langchain.memory import ConversationBufferMemory
from langchain.prompts import PromptTemplate
from langchain_groq import ChatGroq
#from langchain_community.embeddings import HuggingFaceEmbeddings


# Load environment variables
load_dotenv()
GROQ_API_KEY = os.getenv("GROQ_API_KEY")
SERPER_API_KEY = os.getenv("SERPER_API_KEY")

# Initialize clients and models
groq_client = Groq(api_key=GROQ_API_KEY)

# Web search settings
MAX_SEARCH_RESULTS = 5

# In-memory storage for chat history
chat_history = {}

# RAG and LangChain setup
class RAGSystem:
    def __init__(self):
        self.vectorstore = None
        self.embeddings = HuggingFaceEmbeddings(
            model_name="sentence-transformers/all-MiniLM-L6-v2",
            model_kwargs={'device': 'cpu'},
            encode_kwargs={'normalize_embeddings': False}
        )
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            length_function=len,
        )
        self.qa_chain = None
        self.memory = ConversationBufferMemory(
            memory_key="chat_history",
            return_messages=True
        )
        self.current_document = None
        
    def initialize_llm(self):
        """Initialize the LangChain Groq LLM"""
        return ChatGroq(
            groq_api_key=GROQ_API_KEY,
            model_name="llama-3.3-70b-versatile",
            temperature=0.2
        )
    
    def create_vectorstore(self, texts, source_name="document"):
        """Create a vector store from texts"""
        try:
            documents = [Document(page_content=text, metadata={"source": source_name}) for text in texts]
            docs = self.text_splitter.split_documents(documents)
            logger.info(f"Created {len(docs)} document chunks from source: {source_name}")
            
            self.vectorstore = Chroma.from_documents(
                documents=docs, 
                embedding=self.embeddings,
                persist_directory="./chroma_db"
            )
            self.current_document = source_name
            return self.vectorstore
        except Exception as e:
            logger.error(f"Error creating vectorstore: {str(e)}")
            raise
    
    def create_qa_chain(self):
        """Create a QA chain with retrieval"""
        if not self.vectorstore:
            logger.warning("No vectorstore available to create QA chain")
            return None
            
        try:
            llm = self.initialize_llm()
            
            # Create custom prompt that emphasizes using the provided context
            template = """Use the following pieces of context to answer the question at the end. 
            If you don't know the answer based on the context, just say that you don't know. 
            Don't try to make up an answer. Use the context below to answer the question.

            Context: {context}

            Question: {question}

            Answer based on the context:"""
            
            QA_PROMPT = PromptTemplate(
                template=template, 
                input_variables=["context", "question"]
            )
            
            # Create retrieval chain
            retriever = self.vectorstore.as_retriever(
                search_type="similarity",
                search_kwargs={"k": 4}
            )
            
            self.qa_chain = RetrievalQA.from_chain_type(
                llm=llm,
                chain_type="stuff",
                retriever=retriever,
                chain_type_kwargs={
                    "prompt": QA_PROMPT,
                    "memory": self.memory
                },
                return_source_documents=True
            )
            
            logger.info("QA chain created successfully")
            return self.qa_chain
            
        except Exception as e:
            logger.error(f"Error creating QA chain: {str(e)}")
            return None
    
    def query_rag(self, question):
        """Query the RAG system"""
        if not self.qa_chain:
            logger.warning("RAG system not initialized - using fallback response")
            return {"result": "I don't have any documents to reference. Please upload a document or provide a URL first.", "source_documents": []}
        
        try:
            result = self.qa_chain({"query": question})
            logger.info(f"RAG query completed. Found {len(result.get('source_documents', []))} relevant documents")
            return result
        except Exception as e:
            logger.error(f"Error querying RAG system: {str(e)}")
            return {"result": f"Error processing your question: {str(e)}", "source_documents": []}

# Initialize RAG system
rag_system = RAGSystem()

# MCP Servers configuration
MCP_SERVERS = {
    "web_search": {
        "enabled": True,
        "name": "Web Search",
        "description": "Perform real-time web searches"
    },
    "calculator": {
        "enabled": True,
        "name": "Calculator",
        "description": "Perform mathematical calculations"
    },
    "time_date": {
        "enabled": True,
        "name": "Time & Date",
        "description": "Get current time and date information"
    },
    "json_processor": {
        "enabled": True,
        "name": "JSON Processor",
        "description": "Process and analyze JSON data"
    },
    "rag_processor": {
        "enabled": True,
        "name": "RAG System",
        "description": "Document retrieval and analysis"
    }
}

def clean_text(text: str) -> str:
    """Clean and normalize text content"""
    if not text:
        return ""
    text = re.sub(r'\s+', ' ', text).strip()
    return ''.join(char for char in text if char.isprintable())

def process_uploaded_file(file_path: str, is_url: bool = False) -> str:
    """Process uploaded file or URL and return text content"""
    try:
        if is_url:
            # Process URL
            logger.info(f"Processing URL: {file_path}")
            return fetch_web_content(file_path)
        else:
            # Process file based on extension
            ext = os.path.splitext(file_path)[1].lower()
            logger.info(f"Processing file: {file_path} with extension: {ext}")
            
            if ext == '.txt':
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = clean_text(f.read())
                    logger.info(f"Read {len(content)} characters from text file")
                    return content
                    
            elif ext == '.pdf':
                text = ""
                with open(file_path, 'rb') as f:
                    pdf_reader = PdfReader(f)
                    for page_num, page in enumerate(pdf_reader.pages):
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                    logger.info(f"Read {len(text)} characters from PDF with {len(pdf_reader.pages)} pages")
                return clean_text(text)
                
            elif ext == '.csv':
                text = ""
                with open(file_path, 'r', encoding='utf-8') as f:
                    csv_reader = csv.reader(f)
                    for row_num, row in enumerate(csv_reader):
                        if row_num == 0:  # Header row
                            text += "Headers: " + ", ".join(row) + "\n"
                        else:
                            text += f"Row {row_num}: " + ", ".join(row) + "\n"
                    logger.info(f"Read {row_num} rows from CSV file")
                return clean_text(text)
                
            elif ext in ['.doc', '.docx']:
                text = ""
                doc = docx.Document(file_path)
                for para in doc.paragraphs:
                    if para.text.strip():
                        text += para.text + "\n"
                logger.info(f"Read {len(text)} characters from Word document")
                return clean_text(text)
                
            elif ext == '.json':
                with open(file_path, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                    content = clean_text(json.dumps(data, indent=2))
                    logger.info(f"Read JSON data with {len(content)} characters")
                    return content
                    
            else:
                logger.warning(f"Unsupported file extension: {ext}")
                return ""
                
    except Exception as e:
        logger.error(f"Error processing file {file_path}: {str(e)}")
        return ""

def perform_web_search(query: str, max_results: int = MAX_SEARCH_RESULTS) -> list:
    """Perform web search using Serper API and return results"""
    try:
        if not SERPER_API_KEY:
            logger.warning("Serper API key not found. Please set SERPER_API_KEY in your .env file.")
            return []
        
        url = "https://google.serper.dev/search"
        payload = {
            "q": query,
            "num": max_results
        }
        headers = {
            'X-API-KEY': SERPER_API_KEY,
            'Content-Type': 'application/json'
        }
        
        logger.info(f"Performing web search for: {query}")
        response = requests.post(url, headers=headers, json=payload)
        response.raise_for_status()
        data = response.json()
        
        # Extract organic results
        results = []
        if 'organic' in data:
            for item in data['organic']:
                results.append({
                    'title': item.get('title', ''),
                    'href': item.get('link', ''),
                    'snippet': item.get('snippet', '')
                })
        
        logger.info(f"Web search found {len(results)} results")
        return results
        
    except Exception as e:
        logger.error(f"Web search error: {str(e)}")
        return []

def fetch_web_content(url: str) -> str:
    """Fetch and clean content from a URL"""
    try:
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
        }
        logger.info(f"Fetching content from URL: {url}")
        response = requests.get(url, headers=headers, timeout=15)
        response.raise_for_status()
        
        # Simple HTML content extraction
        soup = BeautifulSoup(response.text, 'html.parser')
        
        # Remove script and style elements
        for script in soup(["script", "style", "nav", "header", "footer", "aside"]):
            script.decompose()
        
        # Get text from main content areas
        main_content = soup.find('main') or soup.find('article') or soup.find('div', class_=re.compile(r'content|main|post', re.I))
        
        if main_content:
            text = main_content.get_text()
        else:
            text = soup.get_text()
        
        # Clean text
        lines = (line.strip() for line in text.splitlines())
        chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
        text = ' '.join(chunk for chunk in chunks if chunk)
        
        logger.info(f"Fetched {len(text)} characters from URL")
        return clean_text(text[:5000])  # Limit content length
        
    except Exception as e:
        logger.error(f"Error fetching {url}: {str(e)}")
        return ""

def get_domain_name(url: str) -> str:
    """Extract domain name from URL"""
    try:
        parsed_uri = urlparse(url)
        domain = parsed_uri.netloc
        if domain.startswith('www.'):
            domain = domain[4:]
        return domain
    except:
        return url

def clear_chat_history(session_id: str):
    """Clear chat history for a specific session"""
    global chat_history
    if session_id in chat_history:
        del chat_history[session_id]
    return True

# MCP Server Functions
def mcp_calculator(expression: str) -> str:
    """MCP Calculator - Evaluate mathematical expressions"""
    try:
        # Safe evaluation of mathematical expressions
        allowed_chars = set('0123456789+-*/(). ')
        if not all(c in allowed_chars for c in expression):
            return "Error: Invalid characters in expression"
        
        # Use numpy for safe evaluation
        result = eval(expression, {"__builtins__": None}, {
            "sin": np.sin, "cos": np.cos, "tan": np.tan,
            "sqrt": np.sqrt, "log": np.log, "exp": np.exp,
            "pi": np.pi, "e": np.e
        })
        return f"Result: {result}"
    except Exception as e:
        return f"Error calculating expression: {str(e)}"

def mcp_time_date() -> str:
    """MCP Time & Date - Get current time and date"""
    from datetime import datetime
    now = datetime.now()
    return f"Current date and time: {now.strftime('%Y-%m-%d %H:%M:%S %Z')}"

def mcp_json_process(json_data: str, operation: str = "analyze") -> str:
    """MCP JSON Processor - Analyze or manipulate JSON data"""
    try:
        data = json.loads(json_data)
        
        if operation == "analyze":
            if isinstance(data, dict):
                keys = list(data.keys())
                return f"JSON object with keys: {keys}"
            elif isinstance(data, list):
                return f"JSON array with {len(data)} items"
            else:
                return f"JSON value: {type(data).__name__}"
        
        elif operation == "pretty":
            return json.dumps(data, indent=2)
            
        else:
            return "Available operations: analyze, pretty"
            
    except Exception as e:
        return f"Error processing JSON: {str(e)}"

def mcp_rag_process(query: str, context: str = "") -> str:
    """MCP RAG Processor - Use LangChain RAG for document analysis"""
    try:
        # If we have specific context, add it to the vector store
        if context:
            logger.info(f"Processing RAG query with context length: {len(context)}")
            rag_system.create_vectorstore([context], "user_context")
            rag_system.create_qa_chain()
        
        # Query the RAG system
        result = rag_system.query_rag(query)
        
        # Format the response
        response = result["result"]
        
        # Add source information if available
        if "source_documents" in result and result["source_documents"]:
            sources = list(set([doc.metadata.get("source", "Unknown") for doc in result["source_documents"]]))
            if sources and sources[0] != "Unknown":
                response += f"\n\nSources: {', '.join(sources)}"
        
        logger.info(f"RAG response generated: {len(response)} characters")
        return response
        
    except Exception as e:
        logger.error(f"Error in RAG processing: {str(e)}")
        return f"Error in RAG processing: {str(e)}"

def execute_mcp_tools(query: str, context: str = "") -> List[Dict[str, Any]]:
    """Execute MCP tools based on query content"""
    results = []
    
    # Calculator MCP
    if any(word in query.lower() for word in ["calculate", "math", "equation", "solve", "+", "-", "*", "/"]):
        try:
            # Extract mathematical expression
            import re
            math_pattern = r"(\d+\.?\d*[\s*+\-/\\*^()\s\d+\.?\d*]*)"
            matches = re.findall(math_pattern, query)
            if matches:
                expression = matches[0].strip()
                result = mcp_calculator(expression)
                results.append({
                    "tool": "calculator",
                    "result": result,
                    "expression": expression
                })
        except Exception as e:
            results.append({
                "tool": "calculator",
                "result": f"Error: {str(e)}",
                "expression": "N/A"
            })
    
    # Time & Date MCP
    if any(word in query.lower() for word in ["time", "date", "now", "current", "today", "what day"]):
        result = mcp_time_date()
        results.append({
            "tool": "time_date",
            "result": result
        })
    
    # JSON Processor MCP
    if "json" in query.lower():
        # Look for JSON data in the query
        json_pattern = r"(\{.*\}|\[.*\])"
        matches = re.findall(json_pattern, query, re.DOTALL)
        if matches:
            json_data = matches[0]
            result = mcp_json_process(json_data)
            results.append({
                "tool": "json_processor",
                "result": result,
                "data_preview": json_data[:100] + "..." if len(json_data) > 100 else json_data
            })
    
    # RAG Processor MCP - for document analysis
    if any(word in query.lower() for word in ["document", "pdf", "text", "analyze", "summarize", "context"]) or rag_system.current_document:
        result = mcp_rag_process(query, context)
        results.append({
            "tool": "rag_processor",
            "result": result
        })
    
    return results

def ask_groq_stream(session_id: str, question: str, web_results: list = None, use_web_search: bool = True, context: str = ""):
    """Stream Groq model response with optional web results, MCP tools, RAG, and chat history"""
    history_text = ""
    if session_id in chat_history:
        for q, a in chat_history[session_id][-5:]:
            history_text += f"Q: {q}\nA: {a}\n\n"
    
    # Execute MCP tools (including RAG)
    mcp_results = execute_mcp_tools(question, context)
    mcp_context = ""
    if mcp_results:
        mcp_context = "\nMCP TOOL RESULTS:\n"
        for i, result in enumerate(mcp_results):
            mcp_context += f"[Tool {i+1}: {result['tool']}]\n{result['result']}\n\n"
    
    # Prepare web context if available and web search is enabled
    web_context = ""
    sources = set()
    
    if use_web_search and web_results and len(web_results) > 0:
        web_context = "\nWEB SEARCH RESULTS:\n"
        for i, result in enumerate(web_results):
            if i >= MAX_SEARCH_RESULTS:
                break
                
            # Use the snippet from search results
            content = result.get('snippet', '')
            if not content:
                content = fetch_web_content(result.get('href', ''))
            
            if content:
                domain = get_domain_name(result.get('href', ''))
                title = result.get('title', 'No title')
                web_context += f"[Source {i+1}: {title} - {domain}]\n{content}\n\n"
                sources.add(domain)
    
    # Different prompts based on whether web search is enabled
    if use_web_search:
        prompt = f"""You are an AI assistant that answers questions using web search results, MCP tools, and RAG when available.
You must follow these rules:
1. Use information from the WEB SEARCH RESULTS and MCP TOOL RESULTS below when relevant
2. If conversation starts with Hi, Hello, or Hey, respond with a friendly greeting
3. If conversation starts with What, Where, When, Who, or Why, respond with a detailed answer
4. If using web search results, cite your sources with numbers in brackets like [1], [2], etc.
5. If using MCP tools, mention that you used specialized tools for the response
6. If the answer isn't in the web results, use your general knowledge
7. Be concise but helpful
8. Always provide accurate information

{mcp_context}

{web_context}

CHAT HISTORY:
{history_text}

QUESTION:
{question}

ANSWER:"""
    else:
        prompt = f"""You are an AI assistant that answers questions based on your general knowledge, MCP tools, and RAG.
You must follow these rules:
1. Use your general knowledge, MCP TOOL RESULTS, and RAG to answer questions
2. If conversation starts with Hi, Hello, or Hey, respond with a friendly greeting
3. If conversation starts with What, Where, When, Who, or Why, respond with a detailed answer
4. If using MCP tools, mention that you used specialized tools for the response
5. Be concise but helpful
6. Always provide accurate information

{mcp_context}

CHAT HISTORY:
{history_text}

QUESTION:
{question}

ANSWER:"""
    
    try:
        full_response = ""
        stream = groq_client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.2,
            max_tokens=1024,
            stream=True
        )
        
        for chunk in stream:
            if chunk.choices[0].delta.content:
                token = chunk.choices[0].delta.content
                full_response += token
                yield {"token": token, "complete": False}
        
        # Add source references if we have them and web search is enabled
        if use_web_search and sources:
            source_text = f"\n\nSources: {', '.join(sources)}"
            full_response += source_text
            yield {"token": source_text, "complete": False}
        
        # Add MCP tool references if we used them
        if mcp_results:
            tools_used = ", ".join([result['tool'] for result in mcp_results])
            tools_text = f"\n\nTools used: {tools_used}"
            full_response += tools_text
            yield {"token": tools_text, "complete": False}
        
        # Save to chat history
        if session_id not in chat_history:
            chat_history[session_id] = []
        chat_history[session_id].append((question, full_response))
        
        yield {"complete": True}
    
    except Exception as e:
        yield {"error": str(e), "complete": True}